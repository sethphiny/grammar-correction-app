"""
Document parsing service for extracting text from Word documents
"""

import io
import re
from typing import List, Dict, Any, Optional
from docx import Document
from models.schemas import DocumentData, DocumentLine

# Try to import textract, fallback to alternative if not available
try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False
    print("Warning: textract not available. .doc file support will be limited.")

# Alternative for .doc files
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

def sanitize_text(text: str) -> str:
    """
    Cleans up text extracted from .docx before analysis.
    - Removes control characters and weird XML leftovers.
    - Normalizes whitespace and quotes.
    - Keeps punctuation and letters intact.
    """
    if not text:
        return ""

    # 1️⃣ Encode-decode to remove non-printable or invalid chars
    text = text.encode("utf-8", errors="ignore").decode("utf-8")

    # 2️⃣ Remove weird zero-width spaces, control chars, etc.
    text = re.sub(r"[\u200B-\u200F\u202A-\u202E\u2060-\u206F]", "", text)

    # 3️⃣ Replace multiple newlines/tabs with a single space
    text = re.sub(r"[\r\n\t]+", " ", text)

    # 4️⃣ Collapse multiple spaces
    text = re.sub(r"\s{2,}", " ", text).strip()

    # 5️⃣ Fix spaced contractions (e.g., "It 's" → "It's")
    # Handle common contractions that might have spaces before apostrophes
    spaced_contraction_patterns = [
        (r"\b([Ii]t)\s+('s)\b", r"\1\2"),
        (r"\b([Hh]e)\s+('s)\b", r"\1\2"),
        (r"\b([Ss]he)\s+('s)\b", r"\1\2"),
        (r"\b([Ww]e)\s+('re)\b", r"\1\2"),
        (r"\b([Tt]hey)\s+('re)\b", r"\1\2"),
        (r"\b([Yy]ou)\s+('re)\b", r"\1\2"),
        (r"\b([Ii])\s+('m)\b", r"\1\2"),
        (r"\b([Ii])\s+('ve)\b", r"\1\2"),
        (r"\b([Ii])\s+('ll)\b", r"\1\2"),
        (r"\b([Ii])\s+('d)\b", r"\1\2"),
        (r"\b([Ww]e)\s+('ve)\b", r"\1\2"),
        (r"\b([Ww]e)\s+('ll)\b", r"\1\2"),
        (r"\b([Ww]e)\s+('d)\b", r"\1\2"),
        (r"\b([Tt]hey)\s+('ve)\b", r"\1\2"),
        (r"\b([Tt]hey)\s+('ll)\b", r"\1\2"),
        (r"\b([Tt]hey)\s+('d)\b", r"\1\2"),
        (r"\b([Yy]ou)\s+('ve)\b", r"\1\2"),
        (r"\b([Yy]ou)\s+('ll)\b", r"\1\2"),
        (r"\b([Yy]ou)\s+('d)\b", r"\1\2"),
        (r"\b([Hh]e)\s+('ll)\b", r"\1\2"),
        (r"\b([Hh]e)\s+('d)\b", r"\1\2"),
        (r"\b([Ss]he)\s+('ll)\b", r"\1\2"),
        (r"\b([Ss]he)\s+('d)\b", r"\1\2"),
        (r"\b([Ii]t)\s+('ll)\b", r"\1\2"),
        (r"\b([Ii]t)\s+('d)\b", r"\1\2"),
        (r"\b([Tt]his)\s+('s)\b", r"\1\2"),
        (r"\b([Tt]hat)\s+('s)\b", r"\1\2"),
        (r"\b([Ww]hat)\s+('s)\b", r"\1\2"),
        (r"\b([Ww]ho)\s+('s)\b", r"\1\2"),
        (r"\b([Ww]here)\s+('s)\b", r"\1\2"),
        (r"\b([Ww]hen)\s+('s)\b", r"\1\2"),
        (r"\b([Ww]hy)\s+('s)\b", r"\1\2"),
        (r"\b([Hh]ow)\s+('s)\b", r"\1\2"),
        (r"\b([Tt]here)\s+('s)\b", r"\1\2"),
        (r"\b([Hh]ere)\s+('s)\b", r"\1\2"),
    ]
    
    for pattern, replacement in spaced_contraction_patterns:
        text = re.sub(pattern, replacement, text)

    # 6️⃣ Optionally normalize curly quotes and dashes
    text = text.replace(""", '"').replace(""", '"')
    text = text.replace("'", "'").replace("'", "'")
    text = text.replace("–", "-").replace("—", "-")

    return text

class DocumentParser:
    """Service for parsing Word documents and extracting text with line information"""
    
    def __init__(self):
        self.supported_formats = ['.doc', '.docx']
    
    async def parse_document(self, file_content: bytes, filename: str) -> Optional[DocumentData]:
        """
        Parse a document and extract text with line-by-line information
        
        Args:
            file_content: Raw file content as bytes
            file_content: Original filename
            
        Returns:
            DocumentData object with parsed content
        """
        try:
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'docx':
                return await self._parse_docx(file_content, filename)
            elif file_extension == 'doc':
                return await self._parse_doc(file_content, filename)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            print(f"Error parsing document {filename}: {str(e)}")
            return None
    
    async def _parse_docx(self, file_content: bytes, filename: str) -> DocumentData:
        """Parse DOCX file using python-docx"""
        try:
            # Create file-like object from bytes
            doc_stream = io.BytesIO(file_content)
            doc = Document(doc_stream)
            
            lines = []
            line_number = 1
            total_sentences = 0
            
            # Extract text from paragraphs - preserve all lines including empty ones
            for paragraph in doc.paragraphs:
                # Split paragraph into lines (handle manual line breaks)
                paragraph_lines = paragraph.text.split('\n')
                
                for line_text in paragraph_lines:
                    # Sanitize the text before processing
                    sanitized_text = sanitize_text(line_text.strip())
                    
                    # Add line even if empty to preserve line numbers
                    if sanitized_text:
                        # Split line into sentences
                        sentences = self._split_into_sentences(sanitized_text)
                        
                        lines.append(DocumentLine(
                            line_number=line_number,
                            content=sanitized_text,
                            sentences=sentences
                        ))
                        
                        total_sentences += len(sentences)
                    else:
                        # Add empty line to preserve line numbering
                        lines.append(DocumentLine(
                            line_number=line_number,
                            content="",
                            sentences=[]
                        ))
                    
                    line_number += 1
            
            # Extract metadata
            metadata = {
                'filename': filename,
                'format': 'docx',
                'paragraph_count': len(doc.paragraphs),
                'has_tables': len(doc.tables) > 0,
                'has_images': len(doc.inline_shapes) > 0
            }
            
            return DocumentData(
                filename=filename,
                lines=lines,
                total_lines=len(lines),
                total_sentences=total_sentences,
                metadata=metadata
            )
            
        except Exception as e:
            raise Exception(f"Failed to parse DOCX file: {str(e)}")
    
    async def _parse_doc(self, file_content: bytes, filename: str) -> DocumentData:
        """Parse DOC file using available methods"""
        try:
            import tempfile
            import os
            
            # Save content to temporary file
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                text = None
                extraction_method = None
                
                # Try textract first if available
                if TEXTRACT_AVAILABLE:
                    try:
                        text = textract.process(temp_file_path).decode('utf-8')
                        extraction_method = 'textract'
                    except Exception as e:
                        print(f"Textract failed: {e}")
                
                # Fallback to docx2txt if textract failed or unavailable
                if not text and DOCX2TXT_AVAILABLE:
                    try:
                        text = docx2txt.process(temp_file_path)
                        extraction_method = 'docx2txt'
                    except Exception as e:
                        print(f"docx2txt failed: {e}")
                
                # If both methods failed, try to read as plain text
                if not text:
                    try:
                        with open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            text = f.read()
                        extraction_method = 'plain_text'
                    except Exception as e:
                        print(f"Plain text reading failed: {e}")
                
                if not text:
                    raise Exception("All parsing methods failed. Please convert your .doc file to .docx format.")
                
                # Split into lines - preserve all lines including empty ones
                lines = []
                line_number = 1
                total_sentences = 0
                
                for line_text in text.split('\n'):
                    # Sanitize the text before processing
                    sanitized_text = sanitize_text(line_text.strip())
                    
                    # Add line even if empty to preserve line numbers
                    if sanitized_text:
                        # Split line into sentences
                        sentences = self._split_into_sentences(sanitized_text)
                        
                        lines.append(DocumentLine(
                            line_number=line_number,
                            content=sanitized_text,
                            sentences=sentences
                        ))
                        
                        total_sentences += len(sentences)
                    else:
                        # Add empty line to preserve line numbering
                        lines.append(DocumentLine(
                            line_number=line_number,
                            content="",
                            sentences=[]
                        ))
                    
                    line_number += 1
                
                # Extract metadata
                metadata = {
                    'filename': filename,
                    'format': 'doc',
                    'extraction_method': extraction_method,
                    'fallback_used': extraction_method != 'textract'
                }
                
                return DocumentData(
                    filename=filename,
                    lines=lines,
                    total_lines=len(lines),
                    total_sentences=total_sentences,
                    metadata=metadata
                )
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            raise Exception(f"Failed to parse DOC file: {str(e)}")
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences, handling common abbreviations and edge cases
        
        Args:
            text: Text to split into sentences
            
        Returns:
            List of sentences
        """
        if not text.strip():
            return []
        
        # Common abbreviations that shouldn't end sentences
        abbreviations = {
            'mr.', 'mrs.', 'ms.', 'dr.', 'prof.', 'sr.', 'jr.', 'inc.', 'ltd.', 'corp.',
            'co.', 'st.', 'ave.', 'blvd.', 'rd.', 'etc.', 'vs.', 'e.g.', 'i.e.', 'a.m.',
            'p.m.', 'u.s.', 'u.k.', 'ph.d.', 'm.d.', 'b.a.', 'm.a.', 'b.s.', 'm.s.'
        }
        
        # Use a more robust approach: find sentence boundaries manually
        # This avoids regex issues and handles abbreviations properly
        sentences = []
        current_sentence = ""
        i = 0
        
        while i < len(text):
            char = text[i]
            current_sentence += char
            
            # Check if we've hit a potential sentence ending
            if char in '.!?':
                # Look ahead to see what comes next
                next_chars = text[i+1:i+3].strip()
                
                # If next character is uppercase or we're at the end, this might be a sentence end
                if not next_chars or (len(next_chars) > 0 and next_chars[0].isupper()):
                    # Check if current sentence ends with an abbreviation
                    sentence_lower = current_sentence.lower().strip()
                    ends_with_abbreviation = any(sentence_lower.endswith(abbr) for abbr in abbreviations)
                    
                    # If it doesn't end with an abbreviation, this is a sentence boundary
                    if not ends_with_abbreviation:
                        sentences.append(current_sentence.strip())
                        current_sentence = ""
            
            i += 1
        
        # Add the last sentence if there's any remaining text
        if current_sentence.strip():
            sentences.append(current_sentence.strip())
        
        # Filter out empty sentences
        return [s for s in sentences if s.strip()]
    
    def get_line_range_for_sentence(self, lines: List[DocumentLine], sentence_text: str) -> tuple:
        """
        Find the line range for a sentence that might span multiple lines
        
        Args:
            lines: List of document lines
            sentence_text: The sentence to find
            
        Returns:
            Tuple of (start_line, end_line) or (line_number, line_number) if single line
        """
        for i, line in enumerate(lines):
            if sentence_text in line.content:
                # Check if sentence spans multiple lines
                if i + 1 < len(lines):
                    next_line = lines[i + 1]
                    # Simple heuristic: if sentence ends with lowercase and next line starts with lowercase,
                    # it might be a continuation
                    if (sentence_text[-1].islower() and 
                        next_line.content and 
                        next_line.content[0].islower()):
                        return (line.line_number, next_line.line_number)
                
                return (line.line_number, line.line_number)
        
        return (0, 0)
