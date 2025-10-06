"""
Report generation service for creating DOCX and PDF reports
"""

import os
import tempfile
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from models.schemas import GrammarIssue, DocumentData

class ReportGenerator:
    """Service for generating grammar correction reports in DOCX and PDF formats"""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    async def generate_report(
        self,
        issues: List[GrammarIssue],
        document_data: DocumentData,
        output_filename: str,
        output_format: str
    ) -> str:
        """
        Generate a grammar correction report
        
        Args:
            issues: List of grammar issues found
            document_data: Original document data
            output_filename: Base filename for output
            output_format: 'docx' or 'pdf'
            
        Returns:
            Path to the generated report file
        """
        if output_format.lower() == 'docx':
            return await self._generate_docx_report(issues, document_data, output_filename)
        elif output_format.lower() == 'pdf':
            return await self._generate_pdf_report(issues, document_data, output_filename)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    async def _generate_docx_report(
        self,
        issues: List[GrammarIssue],
        document_data: DocumentData,
        output_filename: str
    ) -> str:
        """Generate DOCX report"""
        doc = Document()
        
        # Set up document styles
        self._setup_docx_styles(doc)
        
        # Add header
        self._add_docx_header(doc, document_data, len(issues))
        
        # Add issues
        self._add_docx_issues(doc, issues)
        
        # Add summary
        self._add_docx_summary(doc, issues)
        
        # Save document
        output_path = os.path.join(self.temp_dir, f"{output_filename}.docx")
        doc.save(output_path)
        
        return output_path
    
    async def _generate_pdf_report(
        self,
        issues: List[GrammarIssue],
        document_data: DocumentData,
        output_filename: str
    ) -> str:
        """Generate PDF report"""
        output_path = os.path.join(self.temp_dir, f"{output_filename}.pdf")
        
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=inch,
            leftMargin=inch,
            topMargin=inch,
            bottomMargin=inch
        )
        
        # Set up styles
        styles = self._setup_pdf_styles()
        
        # Build content
        story = []
        
        # Add header
        self._add_pdf_header(story, styles, document_data, len(issues))
        
        # Add issues
        self._add_pdf_issues(story, styles, issues)
        
        # Add summary
        self._add_pdf_summary(story, styles, issues)
        
        # Build PDF
        doc.build(story)
        
        return output_path
    
    def _setup_docx_styles(self, doc: Document):
        """Set up custom styles for DOCX document"""
        styles = doc.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Header style
        header_style = styles.add_style('CustomHeader', WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.name = 'Arial'
        header_style.font.size = Pt(12)
        header_style.font.bold = True
        header_style.paragraph_format.space_before = Pt(12)
        header_style.paragraph_format.space_after = Pt(6)
        
        # Issue style
        issue_style = styles.add_style('CustomIssue', WD_STYLE_TYPE.PARAGRAPH)
        issue_style.font.name = 'Arial'
        issue_style.font.size = Pt(11)
        issue_style.paragraph_format.space_after = Pt(6)
        issue_style.paragraph_format.left_indent = Inches(0.25)
        
        # Summary style
        summary_style = styles.add_style('CustomSummary', WD_STYLE_TYPE.PARAGRAPH)
        summary_style.font.name = 'Arial'
        summary_style.font.size = Pt(11)
        summary_style.paragraph_format.space_after = Pt(6)
    
    def _setup_pdf_styles(self) -> getSampleStyleSheet:
        """Set up custom styles for PDF document"""
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontName='Helvetica-Bold',
            fontSize=16,
            spaceAfter=12,
            alignment=1  # Center alignment
        )
        styles.add(title_style)
        
        # Header style
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12,
            spaceBefore=12,
            spaceAfter=6
        )
        styles.add(header_style)
        
        # Issue style
        issue_style = ParagraphStyle(
            'CustomIssue',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            spaceAfter=6,
            leftIndent=18
        )
        styles.add(issue_style)
        
        # Summary style
        summary_style = ParagraphStyle(
            'CustomSummary',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            spaceAfter=6
        )
        styles.add(summary_style)
        
        return styles
    
    def _add_docx_header(self, doc: Document, document_data: DocumentData, issues_count: int):
        """Add header to DOCX document"""
        # Title
        title = doc.add_paragraph("Grammar Correction Report", style='CustomTitle')
        
        # Document info
        doc.add_paragraph(f"Original Document: {document_data.filename}", style='CustomHeader')
        doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", style='CustomHeader')
        doc.add_paragraph(f"Total Issues Found: {issues_count}", style='CustomHeader')
        doc.add_paragraph(f"Total Lines Analyzed: {document_data.total_lines}", style='CustomHeader')
        doc.add_paragraph(f"Total Sentences Analyzed: {document_data.total_sentences}", style='CustomHeader')
        
        # Add page break
        doc.add_page_break()
    
    def _add_pdf_header(self, story: List, styles: getSampleStyleSheet, document_data: DocumentData, issues_count: int):
        """Add header to PDF document"""
        # Title
        story.append(Paragraph("Grammar Correction Report", styles['CustomTitle']))
        story.append(Spacer(1, 12))
        
        # Document info
        story.append(Paragraph(f"Original Document: {document_data.filename}", styles['CustomHeader']))
        story.append(Paragraph(f"Analysis Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['CustomHeader']))
        story.append(Paragraph(f"Total Issues Found: {issues_count}", styles['CustomHeader']))
        story.append(Paragraph(f"Total Lines Analyzed: {document_data.total_lines}", styles['CustomHeader']))
        story.append(Paragraph(f"Total Sentences Analyzed: {document_data.total_sentences}", styles['CustomHeader']))
        
        story.append(PageBreak())
    
    def _add_docx_issues(self, doc: Document, issues: List[GrammarIssue]):
        """Add issues to DOCX document"""
        if not issues:
            doc.add_paragraph("No grammar issues found in the document.", style='CustomHeader')
            return
        
        doc.add_paragraph("Issues Found", style='CustomHeader')
        
        for issue in issues:
            # Location
            location = f"Line {issue.line_number}"
            if issue.line_range and issue.line_range != f"{issue.line_number}":
                location = f"Line {issue.line_range}"
            location += f", Sentence {issue.sentence_number}"
            
            doc.add_paragraph(location, style='CustomIssue')
            
            # Original text
            doc.add_paragraph(f'"{issue.original_text}"', style='CustomIssue')
            
            # Problem
            doc.add_paragraph(f"• Problem: {issue.problem}", style='CustomIssue')
            
            # Fix
            doc.add_paragraph(f"• Fix: {issue.fix}", style='CustomIssue')
            
            # Corrected sentence (for spelling errors)
            if issue.corrected_sentence:
                doc.add_paragraph(f"• Corrected sentence: \"{issue.corrected_sentence}\"", style='CustomIssue')
            
            # Add spacing
            doc.add_paragraph("", style='CustomIssue')
    
    def _add_pdf_issues(self, story: List, styles: getSampleStyleSheet, issues: List[GrammarIssue]):
        """Add issues to PDF document"""
        if not issues:
            story.append(Paragraph("No grammar issues found in the document.", styles['CustomHeader']))
            return
        
        story.append(Paragraph("Issues Found", styles['CustomHeader']))
        
        for issue in issues:
            # Location
            location = f"Line {issue.line_number}"
            if issue.line_range and issue.line_range != f"{issue.line_number}":
                location = f"Line {issue.line_range}"
            location += f", Sentence {issue.sentence_number}"
            
            story.append(Paragraph(location, styles['CustomIssue']))
            
            # Original text
            story.append(Paragraph(f'"{issue.original_text}"', styles['CustomIssue']))
            
            # Problem
            story.append(Paragraph(f"• Problem: {issue.problem}", styles['CustomIssue']))
            
            # Fix
            story.append(Paragraph(f"• Fix: {issue.fix}", styles['CustomIssue']))
            
            # Corrected sentence (for spelling errors)
            if issue.corrected_sentence:
                story.append(Paragraph(f"• Corrected sentence: \"{issue.corrected_sentence}\"", styles['CustomIssue']))
            
            # Add spacing
            story.append(Spacer(1, 6))
    
    def _add_docx_summary(self, doc: Document, issues: List[GrammarIssue]):
        """Add summary to DOCX document"""
        doc.add_page_break()
        doc.add_paragraph("✅ Summary", style='CustomHeader')
        
        if not issues:
            doc.add_paragraph("No issues found in the document.", style='CustomSummary')
            return
        
        # Count issues by category
        category_counts = {}
        for issue in issues:
            category = issue.category
            if category in category_counts:
                category_counts[category] += 1
            else:
                category_counts[category] = 1
        
        # Add summary items
        for category, count in category_counts.items():
            category_name = self._get_category_display_name(category)
            doc.add_paragraph(f"• {category_name}: {count}", style='CustomSummary')
        
        # Add totals
        doc.add_paragraph(f"• Total issues: {len(issues)}", style='CustomSummary')
        doc.add_paragraph(f"• Lines with issues: {len(set(issue.line_number for issue in issues))}", style='CustomSummary')
        doc.add_paragraph(f"• Sentences with issues: {len(set((issue.line_number, issue.sentence_number) for issue in issues))}", style='CustomSummary')
    
    def _add_pdf_summary(self, story: List, styles: getSampleStyleSheet, issues: List[GrammarIssue]):
        """Add summary to PDF document"""
        story.append(PageBreak())
        story.append(Paragraph("✅ Summary", styles['CustomHeader']))
        
        if not issues:
            story.append(Paragraph("No issues found in the document.", styles['CustomSummary']))
            return
        
        # Count issues by category
        category_counts = {}
        for issue in issues:
            category = issue.category
            if category in category_counts:
                category_counts[category] += 1
            else:
                category_counts[category] = 1
        
        # Add summary items
        for category, count in category_counts.items():
            category_name = self._get_category_display_name(category)
            story.append(Paragraph(f"• {category_name}: {count}", styles['CustomSummary']))
        
        # Add totals
        story.append(Paragraph(f"• Total issues: {len(issues)}", styles['CustomSummary']))
        story.append(Paragraph(f"• Lines with issues: {len(set(issue.line_number for issue in issues))}", styles['CustomSummary']))
        story.append(Paragraph(f"• Sentences with issues: {len(set((issue.line_number, issue.sentence_number) for issue in issues))}", styles['CustomSummary']))
    
    def _get_category_display_name(self, category: str) -> str:
        """Get display name for category"""
        category_names = {
            'tense_consistency': 'Verb tense consistency issues',
            'subject_verb_agreement': 'Subject-verb agreement',
            'punctuation': 'Grammar/punctuation (commas, quotation marks)',
            'awkward_phrasing': 'Awkward phrasing',
            'redundancy': 'Redundancy',
            'style': 'Style and clarity',
            'spelling': 'Spelling errors'
        }
        
        return category_names.get(category, category.replace('_', ' ').title())
